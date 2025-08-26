from bs4 import BeautifulSoup
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_DIRECTION

def handle_paragraph(tag, doc):
    """
    يعالج الفقرات (p) ويدعم النصوص العريضة والمائلة داخلها.
    """
    p = doc.add_paragraph()
    
    # التحقق من اتجاه النص (للعربية)
    if tag.get('dir') == 'rtl':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # المرور على كل جزء داخل الفقرة (نص عادي، أو وسوم أخرى مثل strong, em)
    for child in tag.children:
        if child.name == 'strong' or child.name == 'b':
            p.add_run(child.get_text()).bold = True
        elif child.name == 'em' or child.name == 'i':
            p.add_run(child.get_text()).italic = True
        else:
            # هذا للتعامل مع النصوص العادية
            p.add_run(str(child))


def handle_table(tag, doc):
    """
    يعالج الجداول ويحولها إلى جدول في مستند Word مع دعم اتجاه اليمين لليسار.
    """
    rows_data = []
    # استخراج كل الصفوف من الجدول
    for row in tag.find_all('tr'):
        cols_data = []
        # استخراج كل الخلايا (th أو td) من كل صف
        for cell in row.find_all(['th', 'td']):
            cols_data.append(cell.get_text(strip=True))
        rows_data.append(cols_data)

    if not rows_data:
        return # لا تقم بإنشاء جدول فارغ

    # إنشاء جدول في الوورد بعدد الصفوف والأعمدة المناسب
    num_rows = len(rows_data)
    num_cols = len(rows_data) if num_rows > 0 else 0
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid' # تطبيق نمط أساسي للجدول

    # التحقق من اتجاه الجدول (للعربية)
    if tag.get('dir') == 'rtl':
        table.direction = WD_TABLE_DIRECTION.RTL

    # تعبئة خلايا الجدول بالبيانات
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            table.cell(i, j).text = cell_text
            # جعل الصف الأول (العناوين) عريضاً
            if i == 0:
                table.cell(i, j).paragraphs.runs.font.bold = True
    
    # إضافة مسافة بعد الجدول لتحسين التنسيق
    doc.add_paragraph()


def convert_html_to_docx(html_path, docx_path):
    """
    الوظيفة الرئيسية التي تقرأ ملف HTML وتحوله إلى DOCX.
    """
    # إنشاء مستند وورد جديد
    doc = docx.Document()

    # قراءة محتوى ملف HTML
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    # تحليل المحتوى باستخدام BeautifulSoup
    soup = BeautifulSoup(html_content, 'lxml')

    # المرور على كل الوسوم الرئيسية داخل جسم الـ HTML
    for tag in soup.body.find_all(recursive=False):
        if tag.name == 'h1':
            p = doc.add_heading(tag.get_text(strip=True), level=1)
            if tag.get('dir') == 'rtl':
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        elif tag.name == 'h2':
            p = doc.add_heading(tag.get_text(strip=True), level=2)
            if tag.get('dir') == 'rtl':
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif tag.name == 'p':
            handle_paragraph(tag, doc)

        elif tag.name == 'ul':
            for li in tag.find_all('li'):
                doc.add_paragraph(li.get_text(strip=True), style='List Bullet')
        
        elif tag.name == 'pre':
            code_text = tag.get_text()
            p = doc.add_paragraph()
            # تنسيق خاص للأكواد
            p_format = p.paragraph_format
            p_format.left_indent = Pt(20)
            p_format.space_before = Pt(4)
            p_format.space_after = Pt(4)
            # إضافة خلفية رمادية (shading)
            shading = p_format.shading
            shading.background_color = RGBColor(240, 240, 240) # رمادي فاتح
            
            run = p.add_run(code_text)
            # استخدام خط أحادي المسافة (monospace)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

        elif tag.name == 'table':
            handle_table(tag, doc)

    # حفظ المستند النهائي
    doc.save(docx_path)
    print(f"تم تحويل الملف بنجاح! تم الحفظ في: {docx_path}")


# نقطة بداية تشغيل السكربت
if __name__ == '__main__':
    html_file = 'sample.html'
    docx_file = 'output.docx'
    convert_html_to_docx(html_file, docx_file)